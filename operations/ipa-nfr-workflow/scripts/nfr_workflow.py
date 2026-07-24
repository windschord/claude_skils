#!/usr/bin/env python3
"""IPA非機能要求グレード ワークフロー管理CLI。

ヒアリング結果の登録、顧客レビュー用Excelの出力、顧客指摘の取り込み、
運用設計書（Word）の生成を行う。

依存ライブラリ: pyyaml（データ保存）、openpyxl（Excel）、python-docx（Word）
    pip install pyyaml openpyxl python-docx

サブコマンド:
    init             データファイル（YAML）を初期化する
    validate         データファイルをロードして全制約を検証する（手編集後の確認用）
    hearing-sheet    ヒアリング質問一覧（未登録項目・優先度順）をMarkdownで出力する
    register         ヒアリング結果（選択レベル・値）をJSONから登録する
    status           登録状況・指摘対応状況のサマリーを表示する
    export-excel     顧客レビュー用Excel（優先度色分け・記入欄付き）を出力する
    import-feedback  顧客が記入したExcelから指摘を取り込む
    list-feedback    取り込んだ指摘を一覧表示する（JSON）
    update-feedback  指摘への対応方針・状態を更新する
    check            要件値の整合性を機械的に検証する
    generate-design  テンプレートへ登録値を機械代入して運用設計書Markdownを生成する
    dump             データ内容をJSON/Markdownで出力する
    export-word      Markdownの運用設計書をWord（.docx）に変換し付録を追加する

== データ整合性の担保機構（重要） ==
永続化される正データは人間可読なYAMLファイル（nfr.yaml）ただ1つ。
SQLiteはディスクに一切保存せず、全サブコマンドが起動時に必ず
「YAMLロード → :memory: SQLiteを nfr_schema.sql から構築 → 全行INSERT」
を行う（load_data関数が唯一の入口）。このINSERT時にCHECK/FK/PK/NOT NULL
の全制約が検証されるため、YAMLの変更（手編集含む）は次のコマンド実行時に
必ず最新のSQLite上で検証される。書き込みは検証済みインメモリDBの直列化
＋ラウンドトリップ検証＋原子的renameで行う（save_dataが唯一の出口）。

設計方針: ドキュメント生成と整合性チェックは本スクリプトが機械的に行い、
YAMLを単一の情報源とすることでドキュメント間の矛盾・記載漏れを排除する。
AI（ヒアリング・顧客指摘への対応方針の判断）とは役割を分離している。
"""

import argparse
import contextlib
import csv
import json
import os
import re
import signal
import sqlite3
import sys
from datetime import datetime, timezone, timedelta
from pathlib import Path

JST = timezone(timedelta(hours=9))
PRIORITY_ORDER = {"高": 0, "中": 1, "低": 2}
JUDGEMENTS = ["未確認", "承認", "要修正", "要協議"]
# 顧客提示物（Excel/設計書/Word）で使うfeedback.statusの日本語表記
STATUS_LABELS = {
    "open": "協議中", "accepted": "対応予定", "rejected": "見送り", "reflected": "反映済み",
}
DEFAULT_MASTER = Path(__file__).resolve().parent.parent / "assets" / "master" / "ipa_nfr_items_ja.csv"
DEFAULT_TEMPLATE = (
    Path(__file__).resolve().parent.parent
    / "assets" / "templates" / "design_template_ja.md"
)
SCHEMA_PATH = Path(__file__).resolve().parent / "nfr_schema.sql"

SHEET_GRADE = "非機能要求グレード"
SHEET_EXTRA = "グレード外要求"
SHEET_GUIDE = "記入ガイド"


def now() -> str:
    """現在日時（JST）を YYYY-MM-DD HH:MM:SS 形式で返す。"""
    return datetime.now(JST).strftime("%Y-%m-%d %H:%M:%S")


def die(msg: str) -> None:
    """エラーメッセージを標準エラーへ出力して終了コード1で終了する。"""
    print(f"エラー: {msg}", file=sys.stderr)
    sys.exit(1)


def _import_yaml():
    """pyyamlを読み込む（未導入時は導入方法を案内して終了する）。"""
    try:
        import yaml
    except ImportError:
        die("pyyaml がインストールされていません: pip install pyyaml")
    return yaml


def _strict_yaml_load(text: str):
    """重複キーをエラーとして検出するYAMLロード。"""
    yaml = _import_yaml()

    class _StrictLoader(yaml.SafeLoader):
        """重複キー検出用のSafeLoader派生クラス。"""

    def _construct_mapping(loader, node, deep=False):
        """マッピング構築時に重複キー（同一項目IDの二重定義等）を検出する。"""
        mapping = {}
        for key_node, value_node in node.value:
            key = loader.construct_object(key_node, deep=deep)
            if key in mapping:
                raise yaml.YAMLError(f"重複キーがあります: {key}")
            mapping[key] = loader.construct_object(value_node, deep=deep)
        return mapping

    _StrictLoader.add_constructor(
        yaml.resolver.BaseResolver.DEFAULT_MAPPING_TAG, _construct_mapping
    )
    return yaml.load(text, Loader=_StrictLoader)


def _build_conn(master_path=None) -> sqlite3.Connection:
    """制約スキーマ（nfr_schema.sql）からインメモリSQLiteを構築し項目マスタをロードする。"""
    if not SCHEMA_PATH.exists():
        die(f"スキーマ定義が見つかりません: {SCHEMA_PATH}")
    master = Path(master_path) if master_path else DEFAULT_MASTER
    if not master.exists():
        die(f"項目マスタが見つかりません: {master}")
    conn = sqlite3.connect(":memory:")
    conn.row_factory = sqlite3.Row
    conn.execute("PRAGMA foreign_keys = ON")
    conn.executescript(SCHEMA_PATH.read_text(encoding="utf-8"))
    with master.open(encoding="utf-8") as f:
        for r in csv.DictReader(f):
            conn.execute(
                "INSERT INTO nfr_items (item_id, category, category_name, item_name,"
                " question, metric_hint, priority, priority_reason, duplicate_of)"
                " VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)",
                (
                    r["item_id"], r["category"], r["category_name"], r["item_name"],
                    r.get("question", ""), r.get("metric_hint", ""), r["priority"],
                    r.get("priority_reason", ""), r.get("duplicate_of") or None,
                ),
            )
    return conn


def load_data(data_path: str) -> sqlite3.Connection:
    """唯一のデータ入口: YAMLをロードし、インメモリSQLiteを構築して全制約を検証する。

    YAMLの変更（手編集を含む）は必ずここを通り、その時点の最新内容で
    再構築されたSQLite上でCHECK/FK/PK/NOT NULL制約の検証を受ける。
    違反があればYAMLを一切変更せずエラー終了する。
    """
    yaml = _import_yaml()
    path = Path(data_path)
    if not path.exists():
        die(f"データファイルが見つかりません: {data_path}（先に init を実行してください）")
    try:
        doc = _strict_yaml_load(path.read_text(encoding="utf-8"))
    except yaml.YAMLError as e:
        die(f"YAMLの読み込みに失敗しました（{data_path}）: {e}")
    if not isinstance(doc, dict) or "project" not in doc:
        die(f"データ形式が不正です（projectセクションがありません）: {data_path}")
    fmt = doc.get("format_version")
    if fmt != 1:
        die(
            f"未対応のformat_versionです: {fmt!r}（本スクリプトは format_version: 1 のみ対応。"
            "欠落している場合は追記してください）"
        )

    project = doc.get("project") or {}
    conn = _build_conn(project.get("master_file"))
    try:
        conn.execute(
            "INSERT INTO project (id, system_name, model_system, master_file,"
            " created_at, updated_at) VALUES (1, ?, ?, ?, ?, ?)",
            (
                project.get("system_name"), project.get("model_system"),
                project.get("master_file"),
                project.get("created_at") or now(), project.get("updated_at") or now(),
            ),
        )
    except sqlite3.IntegrityError as e:
        die(f"データ制約違反（project）: {e}")

    selections = doc.get("selections") or {}
    if not isinstance(selections, dict):
        die("データ形式が不正です（selectionsはマッピングで記述してください）")
    for item_id, s in selections.items():
        s = s or {}
        try:
            conn.execute(
                "INSERT INTO selections (item_id, level, value, note,"
                " customer_judgement, updated_at) VALUES (?, ?, ?, ?, ?, ?)",
                (
                    item_id, s.get("level"), s.get("value"), s.get("note"),
                    s.get("customer_judgement") or "未確認",
                    s.get("updated_at") or now(),
                ),
            )
        except sqlite3.IntegrityError as e:
            die(f"データ制約違反（selections {item_id}）: {e}")

    feedback = doc.get("feedback") or []
    if not isinstance(feedback, list):
        die("データ形式が不正です（feedbackはリストで記述してください）")
    for f in feedback:
        f = f or {}
        try:
            conn.execute(
                "INSERT INTO feedback (id, kind, item_id, classification, feedback,"
                " background, requested_priority, judgement, response, status,"
                " imported_at)"
                " VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)",
                (
                    f.get("id"), f.get("kind"), f.get("item_id"),
                    f.get("classification"), f.get("feedback"), f.get("background"),
                    f.get("requested_priority"), f.get("judgement"),
                    f.get("response"), f.get("status") or "open",
                    f.get("imported_at") or now(),
                ),
            )
        except sqlite3.IntegrityError as e:
            die(f"データ制約違反（feedback #{f.get('id')}）: {e}")
    conn.commit()
    return conn


def _serialize(conn) -> dict:
    """インメモリDBの内容を保存用のYAML構造（辞書）へ直列化する。

    item_name / priority は人間の可読性のための参考表示で、保存のたびに
    項目マスタから自動付記される（ロード時は無視される）。
    """
    p = conn.execute("SELECT * FROM project").fetchone()
    items = {
        r["item_id"]: r
        for r in conn.execute("SELECT item_id, item_name, priority FROM nfr_items")
    }
    doc = {
        "format_version": 1,
        "project": {
            "system_name": p["system_name"],
            "model_system": p["model_system"],
            "master_file": p["master_file"],
            "created_at": p["created_at"],
            "updated_at": p["updated_at"],
        },
        "selections": {},
        "feedback": [],
    }
    for r in conn.execute("SELECT * FROM selections ORDER BY item_id"):
        info = items.get(r["item_id"])
        doc["selections"][r["item_id"]] = {
            "item_name": info["item_name"] if info else None,
            "priority": info["priority"] if info else None,
            "level": r["level"],
            "value": r["value"],
            "note": r["note"],
            "customer_judgement": r["customer_judgement"],
            "updated_at": r["updated_at"],
        }
    for r in conn.execute("SELECT * FROM feedback ORDER BY id"):
        info = items.get(r["item_id"]) if r["item_id"] else None
        doc["feedback"].append({
            "id": r["id"],
            "kind": r["kind"],
            "item_id": r["item_id"],
            "item_name": info["item_name"] if info else None,
            "classification": r["classification"],
            "feedback": r["feedback"],
            "background": r["background"],
            "requested_priority": r["requested_priority"],
            "judgement": r["judgement"],
            "response": r["response"],
            "status": r["status"],
            "imported_at": r["imported_at"],
        })
    return doc


def _core_state(conn) -> tuple:
    """ラウンドトリップ検証用にDBの中核データを正規化して返す。"""
    p = tuple(conn.execute(
        "SELECT system_name, model_system, master_file FROM project"
    ).fetchone())
    sel = [tuple(r) for r in conn.execute(
        "SELECT item_id, level, value, note, customer_judgement, updated_at"
        " FROM selections ORDER BY item_id"
    )]
    fb = [tuple(r) for r in conn.execute(
        "SELECT id, kind, item_id, classification, feedback, background,"
        " requested_priority, judgement, response, status, imported_at"
        " FROM feedback ORDER BY id"
    )]
    return (p, sel, fb)


def save_data(conn, data_path: str) -> None:
    """唯一のデータ出口: 検証済みインメモリDBをYAMLへ直列化して原子的に保存する。

    書き込み前に「直列化 → 再ロード（＝全制約の再検証） → 内容一致」の
    ラウンドトリップ検証を行い、一致しない場合は元ファイルを変更せず終了する。
    """
    yaml = _import_yaml()
    conn.execute("UPDATE project SET updated_at = ?", (now(),))
    conn.commit()
    doc = _serialize(conn)
    text = yaml.safe_dump(
        doc, allow_unicode=True, sort_keys=False, width=10**6,
        default_flow_style=False,
    )
    path = Path(data_path)
    tmp = path.with_name(path.name + ".tmp")
    tmp.write_text(text, encoding="utf-8")
    try:
        conn2 = load_data(str(tmp))
        if _core_state(conn) != _core_state(conn2):
            die(
                "ラウンドトリップ検証に失敗しました（直列化前後で内容が一致しません）。"
                "元のデータファイルは変更していません"
            )
        conn2.close()
    except SystemExit:
        tmp.unlink(missing_ok=True)
        raise
    os.replace(tmp, path)


# ---------------------------------------------------------------- init
def cmd_init(args) -> None:
    """データファイル（YAML）を初期化し、プロジェクト情報を登録する。"""
    if Path(args.data).exists() and not args.force:
        die(f"データファイルが既に存在します: {args.data}（上書きする場合は --force を指定）")
    conn = _build_conn(args.master)
    conn.execute(
        "INSERT INTO project (id, system_name, model_system, master_file,"
        " created_at, updated_at) VALUES (1, ?, ?, ?, ?, ?)",
        (args.project, args.model_system, args.master or None, now(), now()),
    )
    conn.commit()
    counts = conn.execute(
        "SELECT priority, COUNT(*) AS n FROM nfr_items GROUP BY priority"
    ).fetchall()
    summary = {r["priority"]: r["n"] for r in counts}
    save_data(conn, args.data)
    print(f"データファイルを初期化しました: {args.data}")
    print(f"  システム名: {args.project} / モデルシステム{args.model_system}")
    total = sum(summary.values())
    print(
        f"  項目マスタ: {total}項目"
        f"（高: {summary.get('高', 0)} / 中: {summary.get('中', 0)} / 低: {summary.get('低', 0)}）"
    )
    conn.close()


# ---------------------------------------------------------------- validate
def cmd_validate(args) -> None:
    """データファイルをロードして全制約を検証する（手編集後の確認用）。"""
    conn = load_data(args.data)
    n_sel = conn.execute("SELECT COUNT(*) AS n FROM selections").fetchone()["n"]
    n_fb = conn.execute("SELECT COUNT(*) AS n FROM feedback").fetchone()["n"]
    p = conn.execute("SELECT * FROM project").fetchone()
    print(f"検証OK: {args.data}（全制約を通過）")
    print(f"  システム名: {p['system_name']} / 選択結果: {n_sel}項目 / 顧客指摘: {n_fb}件")
    conn.close()


# ---------------------------------------------------------------- hearing-sheet
def cmd_hearing_sheet(args) -> None:
    """ヒアリング質問一覧（デフォルトは未登録項目のみ・優先度順）をMarkdownで出力する。"""
    conn = load_data(args.data)
    p = conn.execute("SELECT * FROM project").fetchone()
    sql = (
        "SELECT i.item_id, i.item_name, i.question, i.metric_hint, i.priority,"
        " i.duplicate_of, s.level, s.value"
        " FROM nfr_items i LEFT JOIN selections s ON s.item_id = i.item_id"
    )
    conds, params = [], []
    if not args.all:
        conds.append("s.item_id IS NULL")
    if args.priority:
        conds.append("i.priority = ?")
        params.append(args.priority)
    if conds:
        sql += " WHERE " + " AND ".join(conds)
    sql += (
        " ORDER BY CASE i.priority WHEN '高' THEN 0 WHEN '中' THEN 1 ELSE 2 END,"
        " i.item_id"
    )
    rows = conn.execute(sql, params).fetchall()

    lines = [
        f"# ヒアリングシート: {p['system_name']}（モデルシステム{p['model_system']}）",
        "",
        f"対象: {'全項目' if args.all else '未登録項目のみ'}"
        f"{f'（優先度: {args.priority}）' if args.priority else ''} / {len(rows)}項目",
        "",
    ]
    for pr in ("高", "中", "低"):
        subset = [r for r in rows if r["priority"] == pr]
        if not subset:
            continue
        lines.append(f"## 優先度: {pr}（{len(subset)}項目）")
        lines.append("")
        header = "| 項目ID | 項目名 | ヒアリング質問 | 記入例 | 備考 |"
        sep = "|--------|--------|---------------|--------|------|"
        if args.all:
            header += " 現在値 |"
            sep += "--------|"
        lines.append(header)
        lines.append(sep)
        for r in subset:
            dup = f"{r['duplicate_of']}と重複（同値を登録）" if r["duplicate_of"] else "—"
            row = (
                f"| {r['item_id']} | {r['item_name']} | {r['question'] or '—'} |"
                f" {r['metric_hint'] or '—'} | {dup} |"
            )
            if args.all:
                current = r["value"] or r["level"] or "未登録"
                row += f" {current} |"
            lines.append(row)
        lines.append("")
    output = "\n".join(lines)
    if args.output:
        Path(args.output).write_text(output, encoding="utf-8")
        print(f"ヒアリングシートを出力しました: {args.output}（{len(rows)}項目）")
    else:
        print(output)
    conn.close()


# ---------------------------------------------------------------- register
def cmd_register(args) -> None:
    """ヒアリング結果JSONを selections へ登録（UPSERT）する。"""
    conn = load_data(args.data)
    if args.input == "-":
        data = json.load(sys.stdin)
    else:
        data = json.loads(Path(args.input).read_text(encoding="utf-8"))
    selections = data if isinstance(data, list) else data.get("selections", [])
    if not selections:
        die("selections が空です")
    known = {r["item_id"] for r in conn.execute("SELECT item_id FROM nfr_items")}
    registered, skipped = 0, []
    for s in selections:
        item_id = s.get("item_id")
        if item_id not in known:
            skipped.append(item_id)
            continue
        conn.execute(
            "INSERT INTO selections (item_id, level, value, note, updated_at)"
            " VALUES (?, ?, ?, ?, ?)"
            " ON CONFLICT(item_id) DO UPDATE SET"
            " level = excluded.level, value = excluded.value, note = excluded.note,"
            " updated_at = excluded.updated_at",
            (item_id, s.get("level"), s.get("value"), s.get("note"), now()),
        )
        registered += 1
    conn.commit()
    save_data(conn, args.data)
    print(f"登録完了: {registered}項目")
    if skipped:
        print(f"  マスタに存在しないためスキップ: {', '.join(map(str, skipped))}")
    for d in _one_sided_dups(conn):
        registered_side = d["item_id"] if d["self_reg"] else d["duplicate_of"]
        missing_side = d["duplicate_of"] if d["self_reg"] else d["item_id"]
        print(
            f"  注意: {registered_side} は {missing_side} と重複項目です。"
            f"{missing_side} にも同じ値を登録してください"
        )
    _print_missing_high(conn)
    conn.close()


def _one_sided_dups(conn) -> list:
    """重複項目（○マーク）ペアのうち片側だけ登録されているものを返す。"""
    return conn.execute(
        "SELECT i.item_id, i.duplicate_of,"
        " (s1.item_id IS NOT NULL) AS self_reg, (s2.item_id IS NOT NULL) AS rep_reg"
        " FROM nfr_items i"
        " LEFT JOIN selections s1 ON s1.item_id = i.item_id"
        " LEFT JOIN selections s2 ON s2.item_id = i.duplicate_of"
        " WHERE i.duplicate_of IS NOT NULL"
        " AND (s1.item_id IS NULL) != (s2.item_id IS NULL)"
    ).fetchall()


def _print_missing_high(conn) -> None:
    """未登録の優先度「高」項目の一覧を表示する。"""
    missing = conn.execute(
        "SELECT i.item_id, i.item_name FROM nfr_items i"
        " LEFT JOIN selections s ON s.item_id = i.item_id"
        " WHERE i.priority = '高' AND s.item_id IS NULL ORDER BY i.item_id"
    ).fetchall()
    if missing:
        print(f"  未登録の優先度「高」項目: {len(missing)}件")
        for r in missing:
            print(f"    - {r['item_id']} {r['item_name']}")


# ---------------------------------------------------------------- status
def cmd_status(args) -> None:
    """優先度別の登録状況と顧客指摘の対応状況サマリーを表示する。"""
    conn = load_data(args.data)
    p = conn.execute("SELECT * FROM project").fetchone()
    if not p:
        die("プロジェクト情報がありません")
    print(f"システム名: {p['system_name']} / モデルシステム{p['model_system']}")
    rows = conn.execute(
        "SELECT i.priority, COUNT(*) AS total, COUNT(s.item_id) AS registered"
        " FROM nfr_items i LEFT JOIN selections s ON s.item_id = i.item_id"
        " GROUP BY i.priority"
    ).fetchall()
    print("登録状況（優先度別）:")
    for pr in ("高", "中", "低"):
        r = next((x for x in rows if x["priority"] == pr), None)
        if r:
            print(f"  {pr}: {r['registered']}/{r['total']} 項目")
    fb = conn.execute(
        "SELECT status, COUNT(*) AS n FROM feedback GROUP BY status"
    ).fetchall()
    if fb:
        print("顧客指摘の状況:")
        for r in fb:
            print(f"  {r['status']}: {r['n']}件")
    else:
        print("顧客指摘: 未取り込み")
    _print_missing_high(conn)
    conn.close()


# ---------------------------------------------------------------- export-excel
def _select_rows(conn):
    """項目マスタと選択結果を結合し、優先度（高→中→低）・項目ID順で返す。"""
    return conn.execute(
        "SELECT i.item_id, i.category, i.category_name, i.item_name, i.metric_hint,"
        " i.priority, i.duplicate_of, s.level, s.value, s.note, s.customer_judgement"
        " FROM nfr_items i LEFT JOIN selections s ON s.item_id = i.item_id"
        " ORDER BY CASE i.priority WHEN '高' THEN 0 WHEN '中' THEN 1 ELSE 2 END,"
        " i.item_id"
    ).fetchall()


def cmd_export_excel(args) -> None:
    """顧客レビュー用Excel（優先度色分け・顧客記入欄・記入ガイド付き）を出力する。"""
    try:
        from openpyxl import Workbook
        from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
        from openpyxl.utils import get_column_letter
        from openpyxl.worksheet.datavalidation import DataValidation
    except ImportError:
        die("openpyxl がインストールされていません: pip install openpyxl")

    conn = load_data(args.data)
    p = conn.execute("SELECT * FROM project").fetchone()
    rows = _select_rows(conn)

    priority_fill = {
        "高": PatternFill("solid", fgColor="FDE9E9"),
        "中": PatternFill("solid", fgColor="FFF6DD"),
        "低": PatternFill("solid", fgColor="F2F2F2"),
    }
    priority_font = {
        "高": Font(color="B03030", bold=True),
        "中": Font(color="8A6D1A", bold=True),
        "低": Font(color="666666"),
    }
    header_fill = PatternFill("solid", fgColor="2F5496")
    header_font = Font(color="FFFFFF", bold=True)
    customer_fill = PatternFill("solid", fgColor="E7F0E7")
    info_fill = PatternFill("solid", fgColor="EAEFF7")
    thin = Side(style="thin", color="BFBFBF")
    border = Border(left=thin, right=thin, top=thin, bottom=thin)
    wrap = Alignment(wrap_text=True, vertical="top")

    wb = Workbook()

    # --- シート1: 非機能要求グレード ---
    ws = wb.active
    ws.title = SHEET_GRADE
    title = f"非機能要件確認シート - {p['system_name']}（モデルシステム{p['model_system']}）"
    ws.cell(row=1, column=1, value=title).font = Font(bold=True, size=14)
    ws.cell(
        row=2, column=1,
        value="優先度「高」は後続の設計作業への影響が大きい項目です。高→中→低の順にご確認ください。"
        "「顧客判定」「顧客コメント」列（緑色）にご記入ください。"
        "再提示の場合、前回いただいた指摘への対応方針を「前回指摘と対応方針」列に掲載しています。",
    ).font = Font(color="808080")

    # 前回レビューの指摘と対応方針（再提示時に顧客が確認できるよう掲載する）
    fb_by_item = {}
    for f in conn.execute(
        "SELECT item_id, feedback, response, status FROM feedback"
        " WHERE kind = 'item' AND item_id IS NOT NULL ORDER BY id"
    ):
        label = STATUS_LABELS.get(f["status"], f["status"])
        fb_by_item.setdefault(f["item_id"], []).append(
            f"指摘: {f['feedback']}\n→ 対応: {f['response'] or '検討中'}（{label}）"
        )

    headers = [
        ("No", 5), ("優先度", 8), ("カテゴリ", 16), ("項目ID", 10), ("項目名", 30),
        ("メトリクス（記入例）", 24), ("選択レベル", 10), ("選択値・具体値", 24),
        ("備考（設計根拠）", 30), ("前回指摘と対応方針", 40), ("顧客判定", 12),
        ("顧客コメント（指摘・要望）", 40),
    ]
    header_row = 4
    for col, (name, width) in enumerate(headers, start=1):
        c = ws.cell(row=header_row, column=col, value=name)
        c.fill = header_fill
        c.font = header_font
        c.border = border
        c.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
        ws.column_dimensions[get_column_letter(col)].width = width

    dv = DataValidation(
        type="list", formula1='"' + ",".join(JUDGEMENTS) + '"', allow_blank=True,
        showErrorMessage=True, errorTitle="入力エラー",
        error="未確認・承認・要修正・要協議 のいずれかを選択してください",
    )
    ws.add_data_validation(dv)

    for i, r in enumerate(rows, start=1):
        row = header_row + i
        dup = f"（{r['duplicate_of']}と重複項目）" if r["duplicate_of"] else ""
        values = [
            i, r["priority"], f"{r['category']}: {r['category_name']}", r["item_id"],
            r["item_name"] + dup, r["metric_hint"], r["level"] or "",
            r["value"] or "", r["note"] or "",
            "\n".join(fb_by_item.get(r["item_id"], [])),
            r["customer_judgement"] or "未確認", "",
        ]
        for col, v in enumerate(values, start=1):
            c = ws.cell(row=row, column=col, value=v)
            c.border = border
            c.alignment = wrap
            if col in (11, 12):
                c.fill = customer_fill
            elif col == 10:
                c.fill = info_fill
        pc = ws.cell(row=row, column=2)
        pc.fill = priority_fill[r["priority"]]
        pc.font = priority_font[r["priority"]]
        pc.alignment = Alignment(horizontal="center", vertical="top")
        dv.add(ws.cell(row=row, column=11))

    ws.freeze_panes = ws.cell(row=header_row + 1, column=6)
    ws.auto_filter.ref = f"A{header_row}:L{header_row + len(rows)}"

    # --- シート2: グレード外要求 ---
    ws2 = wb.create_sheet(SHEET_EXTRA)
    ws2.cell(
        row=1, column=1,
        value="非機能要求グレードの項目では表現できないご要望・ご指摘は本シートにご記入ください。",
    ).font = Font(bold=True)
    extra_headers = [
        ("No", 5), ("分類", 16), ("要求・指摘内容", 50), ("背景・理由", 40),
        ("希望優先度", 12), ("備考", 24),
    ]
    for col, (name, width) in enumerate(extra_headers, start=1):
        c = ws2.cell(row=3, column=col, value=name)
        c.fill = header_fill
        c.font = header_font
        c.border = border
        c.alignment = Alignment(horizontal="center", vertical="center")
        ws2.column_dimensions[get_column_letter(col)].width = width
    dv2 = DataValidation(type="list", formula1='"高,中,低"', allow_blank=True)
    ws2.add_data_validation(dv2)
    # 取込済みのグレード外要求は対応方針つきで再掲する（再提示時の確認用。編集不要）
    existing_extra = conn.execute(
        "SELECT * FROM feedback WHERE kind = 'out_of_grade' ORDER BY id"
    ).fetchall()
    row = 3
    for i, f in enumerate(existing_extra, start=1):
        row = 3 + i
        label = STATUS_LABELS.get(f["status"], f["status"])
        values2 = [
            i, f["classification"] or "", f["feedback"], f["background"] or "",
            f["requested_priority"] or "",
            f"対応方針: {f['response'] or '検討中'}（{label}）",
        ]
        for col, v in enumerate(values2, start=1):
            c = ws2.cell(row=row, column=col, value=v)
            c.border = border
            c.alignment = wrap
            c.fill = info_fill
    for i in range(1, 21):
        row = 3 + len(existing_extra) + i
        ws2.cell(row=row, column=1, value=len(existing_extra) + i).border = border
        for col in range(2, 7):
            c = ws2.cell(row=row, column=col)
            c.border = border
            c.alignment = wrap
            c.fill = customer_fill
        dv2.add(ws2.cell(row=row, column=5))

    # --- シート3: 記入ガイド ---
    ws3 = wb.create_sheet(SHEET_GUIDE)
    guide = [
        ("本シートの目的", ""),
        ("", "IPA「非機能要求グレード2018」に基づき、対象システムの非機能要件の選択結果をご確認いただくためのシートです。"),
        ("", ""),
        ("優先度の意味（後続の設計・構築作業への影響度）", ""),
        ("高", "インフラ構成・アーキテクチャ・運用体制の根幹を決める項目。変更の後戻りコストが大きいため、最優先でご確認ください。"),
        ("中", "運用プロセス・手順の設計に影響する項目。設計工程の中で調整可能ですが、早期の確定が望まれます。"),
        ("低", "付加的・詳細レベルの項目。後工程での調整が比較的容易です。"),
        ("", ""),
        ("顧客判定の記入方法", ""),
        ("承認", "選択レベル・値のままで問題ない場合"),
        ("要修正", "値の変更が必要な場合（顧客コメント欄に修正内容をご記入ください）"),
        ("要協議", "打ち合わせでの協議を希望される場合（顧客コメント欄に論点をご記入ください）"),
        ("未確認", "まだ確認されていない項目（初期値）"),
        ("", ""),
        ("グレード外要求シートについて", ""),
        ("", "非機能要求グレードの項目体系では表現できないご要望（例: 特定の運用レポート、独自の承認フロー、組織固有の制約など）は「グレード外要求」シートに自由記述でご記入ください。運用設計書に反映します。"),
        ("", ""),
        ("未記入項目への値のご提案", ""),
        ("", "「選択レベル」「選択値・具体値」が空欄の項目に値をご提案いただく場合は、「顧客コメント（指摘・要望）」欄にご記入ください。「選択値・具体値」列へ直接ご記入いただいても取り込まれませんのでご注意ください。"),
        ("", ""),
        ("再提示時の見方", ""),
        ("", "「前回指摘と対応方針」列（薄青色）には、前回のレビューでいただいた指摘と当社の対応方針を掲載しています。グレード外要求シートの薄青色の行も同様に、取込済みのご要望と対応方針の再掲です（編集不要）。"),
    ]
    ws3.column_dimensions["A"].width = 14
    ws3.column_dimensions["B"].width = 100
    for i, (a, b) in enumerate(guide, start=1):
        ca = ws3.cell(row=i, column=1, value=a)
        cb = ws3.cell(row=i, column=2, value=b)
        cb.alignment = wrap
        if a and not b:
            ca.font = Font(bold=True, size=12)
        elif a in priority_fill:
            ca.fill = priority_fill[a]
            ca.font = priority_font[a]
            ca.alignment = Alignment(horizontal="center")

    wb.save(args.output)
    high = sum(1 for r in rows if r["priority"] == "高")
    print(f"Excel出力完了: {args.output}")
    print(f"  {SHEET_GRADE}: {len(rows)}項目（うち優先度「高」{high}項目）")
    print(f"  {SHEET_EXTRA}: 記入欄20行 / {SHEET_GUIDE}: あり")
    conn.close()


# ---------------------------------------------------------------- import-feedback
def _header_map(ws, header_row: int) -> dict:
    """指定行のヘッダー文字列から列番号への対応表を作る。"""
    return {
        (c.value or "").strip(): c.column
        for c in ws[header_row] if isinstance(c.value, str)
    }


def cmd_import_feedback(args) -> None:
    """顧客記入済みExcelから顧客判定と指摘（グレード外要求含む）を取り込む。"""
    try:
        from openpyxl import load_workbook
    except ImportError:
        die("openpyxl がインストールされていません: pip install openpyxl")
    if not Path(args.input).exists():
        die(f"Excelファイルが見つかりません: {args.input}")

    conn = load_data(args.data)
    wb = load_workbook(args.input, data_only=True)
    imported, updated_judgements, skipped_dup = [], 0, 0

    # --- シート1: 項目別の判定・コメント ---
    if SHEET_GRADE not in wb.sheetnames:
        die(f"シート「{SHEET_GRADE}」が見つかりません")
    ws = wb[SHEET_GRADE]
    hmap = _header_map(ws, 4)
    for need in ("項目ID", "顧客判定", "顧客コメント（指摘・要望）"):
        if need not in hmap:
            die(f"シート「{SHEET_GRADE}」に列「{need}」が見つかりません")
    known = {r["item_id"] for r in conn.execute("SELECT item_id FROM nfr_items")}
    for row in ws.iter_rows(min_row=5):
        item_id = row[hmap["項目ID"] - 1].value
        if not item_id or str(item_id).strip() not in known:
            continue
        item_id = str(item_id).strip()
        judgement = str(row[hmap["顧客判定"] - 1].value or "").strip()
        comment = str(row[hmap["顧客コメント（指摘・要望）"] - 1].value or "").strip()
        if judgement and judgement not in JUDGEMENTS:
            print(
                f"  警告: {item_id} の顧客判定「{judgement}」は許容値"
                f"（{'/'.join(JUDGEMENTS)}）外のためスキップします",
                file=sys.stderr,
            )
            judgement = ""
        if judgement and judgement != "未確認":
            cur = conn.execute(
                "UPDATE selections SET customer_judgement = ?, updated_at = ?"
                " WHERE item_id = ?",
                (judgement, now(), item_id),
            )
            if cur.rowcount:
                updated_judgements += 1
            else:
                print(
                    f"  警告: {item_id} は選択結果が未登録のため顧客判定"
                    f"「{judgement}」を反映できません（registerで登録後に再取込してください）",
                    file=sys.stderr,
                )
        if comment:
            exists = conn.execute(
                "SELECT 1 FROM feedback WHERE kind = 'item' AND item_id = ?"
                " AND feedback = ?",
                (item_id, comment),
            ).fetchone()
            if exists:
                skipped_dup += 1
                continue
            cur = conn.execute(
                "INSERT INTO feedback (kind, item_id, feedback, judgement, imported_at)"
                " VALUES ('item', ?, ?, ?, ?)",
                (item_id, comment, judgement or None, now()),
            )
            imported.append((cur.lastrowid, "item", item_id, comment))

    # --- シート2: グレード外要求 ---
    if SHEET_EXTRA in wb.sheetnames:
        ws2 = wb[SHEET_EXTRA]
        hmap2 = _header_map(ws2, 3)
        col_content = hmap2.get("要求・指摘内容")
        if not col_content:
            die(
                f"シート「{SHEET_EXTRA}」に列「要求・指摘内容」が見つかりません"
                "（ヘッダー行が変更されていないか確認してください）"
            )
        if col_content:
            for row in ws2.iter_rows(min_row=4):
                content = str(row[col_content - 1].value or "").strip()
                if not content:
                    continue
                classification = str(
                    row[hmap2.get("分類", 2) - 1].value or ""
                ).strip() or None
                background = str(
                    row[hmap2.get("背景・理由", 4) - 1].value or ""
                ).strip() or None
                req_priority = str(
                    row[hmap2.get("希望優先度", 5) - 1].value or ""
                ).strip() or None
                exists = conn.execute(
                    "SELECT 1 FROM feedback WHERE kind = 'out_of_grade' AND feedback = ?",
                    (content,),
                ).fetchone()
                if exists:
                    skipped_dup += 1
                    continue
                cur = conn.execute(
                    "INSERT INTO feedback (kind, classification, feedback, background,"
                    " requested_priority, imported_at)"
                    " VALUES ('out_of_grade', ?, ?, ?, ?, ?)",
                    (classification, content, background, req_priority, now()),
                )
                imported.append((cur.lastrowid, "out_of_grade", classification, content))

    conn.commit()
    save_data(conn, args.data)
    print(f"取り込み完了: {args.input}")
    print(f"  顧客判定の更新: {updated_judgements}項目")
    print(f"  新規指摘: {len(imported)}件（重複スキップ: {skipped_dup}件）")
    for fid, kind, target, text in imported:
        label = target if kind == "item" else f"グレード外/{target or '分類なし'}"
        print(f"    [#{fid}] {label}: {text[:60]}")
    conn.close()


# ---------------------------------------------------------------- list/update feedback
def cmd_list_feedback(args) -> None:
    """取り込んだ顧客指摘の一覧をJSONで出力する。"""
    conn = load_data(args.data)
    sql = (
        "SELECT f.*, i.item_name, i.priority FROM feedback f"
        " LEFT JOIN nfr_items i ON i.item_id = f.item_id"
    )
    params = []
    if args.status:
        sql += " WHERE f.status = ?"
        params.append(args.status)
    sql += " ORDER BY f.id"
    rows = [dict(r) for r in conn.execute(sql, params)]
    print(json.dumps(rows, ensure_ascii=False, indent=2))
    conn.close()


def cmd_update_feedback(args) -> None:
    """顧客指摘の対応方針・状態・項目紐付けを更新する。"""
    conn = load_data(args.data)
    row = conn.execute("SELECT * FROM feedback WHERE id = ?", (args.id,)).fetchone()
    if not row:
        die(f"指摘 #{args.id} が見つかりません")
    sets, params = [], []
    if args.response is not None:
        sets.append("response = ?")
        params.append(args.response)
    if args.status:
        sets.append("status = ?")
        params.append(args.status)
    if args.item_id:
        known = conn.execute(
            "SELECT 1 FROM nfr_items WHERE item_id = ?", (args.item_id,)
        ).fetchone()
        if not known:
            die(f"項目ID {args.item_id} はマスタに存在しません")
        sets.append("item_id = ?")
        params.append(args.item_id)
    if not sets:
        die("--response / --status / --item-id のいずれかを指定してください")
    params.append(args.id)
    conn.execute(f"UPDATE feedback SET {', '.join(sets)} WHERE id = ?", params)
    conn.commit()
    save_data(conn, args.data)
    print(f"指摘 #{args.id} を更新しました")
    conn.close()


# ---------------------------------------------------------------- dump
def cmd_dump(args) -> None:
    """DB内容（プロジェクト・選択結果・指摘）をJSONまたはMarkdownで出力する。"""
    conn = load_data(args.data)
    p = dict(conn.execute("SELECT * FROM project").fetchone())
    items = [dict(r) for r in _select_rows(conn)]
    feedback = [
        dict(r) for r in conn.execute(
            "SELECT f.*, i.item_name FROM feedback f"
            " LEFT JOIN nfr_items i ON i.item_id = f.item_id ORDER BY f.id"
        )
    ]
    if args.format == "json":
        print(json.dumps(
            {"project": p, "items": items, "feedback": feedback},
            ensure_ascii=False, indent=2,
        ))
    else:
        print(f"# 非機能要件データ: {p['system_name']}（モデルシステム{p['model_system']}）\n")
        for pr in ("高", "中", "低"):
            subset = [i for i in items if i["priority"] == pr]
            print(f"## 優先度: {pr}（{len(subset)}項目）\n")
            print("| 項目ID | 項目名 | 選択レベル | 選択値 | 備考 | 顧客判定 |")
            print("|--------|--------|-----------|--------|------|---------|")
            for i in subset:
                print(
                    f"| {i['item_id']} | {i['item_name']} | {i['level'] or '未登録'} |"
                    f" {i['value'] or ''} | {i['note'] or ''} |"
                    f" {i['customer_judgement'] or '未確認'} |"
                )
            print()
        if feedback:
            print("## 顧客指摘一覧\n")
            print("| # | 区分 | 対象 | 指摘・要求内容 | 対応方針 | 状態 |")
            print("|---|------|------|---------------|---------|------|")
            for f in feedback:
                target = (
                    f"{f['item_id']} {f['item_name'] or ''}".strip()
                    if f["kind"] == "item"
                    else f"グレード外（{f['classification'] or '分類なし'}）"
                )
                print(
                    f"| {f['id']} | {'項目指摘' if f['kind'] == 'item' else 'グレード外要求'} |"
                    f" {target} | {f['feedback']} | {f['response'] or '未定'} | {f['status']} |"
                )
    conn.close()


# ---------------------------------------------------------------- check
DURATION_RE = re.compile(r"(\d+(?:\.\d+)?)\s*(分|時間|日|週間|ヶ月|か月|カ月)")
DURATION_KEYWORDS = {
    "リアルタイム": 0.0, "常時": 0.0, "時間毎": 1.0, "毎時": 1.0,
    "日次": 24.0, "毎日": 24.0, "週次": 168.0, "毎週": 168.0,
    "月次": 720.0, "毎月": 720.0, "年次": 8760.0,
}
UNIT_HOURS = {"分": 1 / 60, "時間": 1.0, "日": 24.0, "週間": 168.0, "ヶ月": 720.0, "か月": 720.0, "カ月": 720.0}
RPO_LEVEL_HOURS = {"L1": 24.0, "L2": 8.0, "L3": 1.0, "L4": 0.25, "L5": 5 / 60}
FULLTIME_RE = re.compile(r"24\s*(時間|h|H)|24/7|365")


def _parse_duration_hours(text: str):
    """自由記述の期間表現から最小の間隔（時間単位）を抽出する。解釈不能ならNone。"""
    if not text:
        return None
    candidates = [
        v for k, v in DURATION_KEYWORDS.items() if k in text
    ] + [
        float(m.group(1)) * UNIT_HOURS[m.group(2)] for m in DURATION_RE.finditer(text)
    ]
    return min(candidates) if candidates else None


def _get_selection(conn, item_id: str):
    """指定項目の選択結果（level/value/note）を返す。未登録ならNone。"""
    return conn.execute(
        "SELECT level, value, note FROM selections WHERE item_id = ?", (item_id,)
    ).fetchone()


def cmd_check(args) -> None:
    """DB内の要件値の整合性を機械的に検証し、ERROR/WARN/UNVERIFIEDを報告する。"""
    conn = load_data(args.data)
    errors, warns, infos = [], [], []

    # 1. 重複項目（○マーク）の値一致
    dup_rows = conn.execute(
        "SELECT i.item_id, i.duplicate_of, s1.value AS v1, s1.level AS l1,"
        " s2.value AS v2, s2.level AS l2"
        " FROM nfr_items i"
        " JOIN selections s1 ON s1.item_id = i.item_id"
        " JOIN selections s2 ON s2.item_id = i.duplicate_of"
        " WHERE i.duplicate_of IS NOT NULL"
    ).fetchall()
    for r in dup_rows:
        v1, v2 = (r["v1"] or "").strip(), (r["v2"] or "").strip()
        l1, l2 = (r["l1"] or "").strip(), (r["l2"] or "").strip()
        if v1 != v2 or (l1 and l2 and l1 != l2):
            errors.append(
                f"重複項目の不一致: {r['item_id']}（{v1 or l1}）と"
                f" {r['duplicate_of']}（{v2 or l2}）の値が異なります"
            )

    # 1b. 重複項目の片側のみ登録（登録漏れ）
    for d in _one_sided_dups(conn):
        registered_side = d["item_id"] if d["self_reg"] else d["duplicate_of"]
        missing_side = d["duplicate_of"] if d["self_reg"] else d["item_id"]
        warns.append(
            f"重複項目の片側未登録: {registered_side} は登録済みですが、重複項目の"
            f" {missing_side} が未登録です（同じ値を登録してください）"
        )

    # 2. バックアップ取得間隔（C.1.2.5） ≦ RPO（A.1.3.1）
    backup = _get_selection(conn, "C.1.2.5")
    rpo = _get_selection(conn, "A.1.3.1")
    if backup and rpo:
        backup_h = _parse_duration_hours(backup["value"] or "")
        rpo_h = _parse_duration_hours(rpo["value"] or "")
        if rpo_h is None and (rpo["level"] or "") in RPO_LEVEL_HOURS:
            rpo_h = RPO_LEVEL_HOURS[rpo["level"]]
        if backup_h is not None and rpo_h is not None:
            if backup_h > rpo_h:
                errors.append(
                    f"バックアップ間隔（C.1.2.5: {backup['value']} ≒ {backup_h:g}時間）が"
                    f" RPO（A.1.3.1: ≒ {rpo_h:g}時間）を超えています"
                )
        else:
            infos.append(
                "UNVERIFIED: バックアップ間隔とRPOの大小関係を機械判定できません"
                f"（C.1.2.5「{backup['value']}」/ A.1.3.1「{rpo['value']}」）。目視確認してください"
            )

    # 3. 稼働率レベル（A.2.1.1）とRTO/RPOレベルの対応（IPA推奨は同一レベル）
    avail = _get_selection(conn, "A.2.1.1")
    if avail and (avail["level"] or "").startswith("L"):
        for target in ("A.1.3.2", "A.1.3.1"):
            sel = _get_selection(conn, target)
            if sel and (sel["level"] or "").startswith("L") and sel["level"] != avail["level"]:
                warns.append(
                    f"稼働率レベル（A.2.1.1: {avail['level']}）と {target}"
                    f"（{sel['level']}）のレベルが一致していません（IPA推奨は同一レベル）"
                )

    # 4. 24時間365日稼働と運用・監視時間帯（C.1.1.2）の整合
    service = _get_selection(conn, "A.1.2.1") or _get_selection(conn, "C.1.1.1")
    monitoring = _get_selection(conn, "C.1.1.2")
    if service and monitoring and FULLTIME_RE.search(service["value"] or ""):
        mon_value = monitoring["value"] or ""
        covered = FULLTIME_RE.search(mon_value) or re.search(
            r"オンコール|自動通報|自動検知", mon_value
        )
        if not covered:
            warns.append(
                f"サービスは24時間稼働（{service['value']}）ですが、運用・監視時間帯"
                f"（C.1.1.2: {monitoring['value']}）が24時間をカバーしていない可能性があります"
            )

    # 5. 優先度「高」の未登録項目
    missing_high = conn.execute(
        "SELECT COUNT(*) AS n FROM nfr_items i"
        " LEFT JOIN selections s ON s.item_id = i.item_id"
        " WHERE i.priority = '高' AND s.item_id IS NULL"
    ).fetchone()["n"]
    if missing_high:
        warns.append(f"優先度「高」の未登録項目が {missing_high}件あります（status で一覧確認）")

    # 6. 未対応（open）の顧客指摘
    open_fb = conn.execute(
        "SELECT COUNT(*) AS n FROM feedback WHERE status = 'open'"
    ).fetchone()["n"]
    if open_fb:
        warns.append(f"未対応（open）の顧客指摘が {open_fb}件あります（list-feedback --status open で確認）")

    for msg in errors:
        print(f"[ERROR] {msg}")
    for msg in warns:
        print(f"[WARN]  {msg}")
    for msg in infos:
        print(f"[INFO]  {msg}")
    print(
        f"整合性チェック完了: ERROR {len(errors)}件 / WARN {len(warns)}件"
        f" / INFO {len(infos)}件"
    )
    conn.close()
    if errors:
        sys.exit(1)


# ---------------------------------------------------------------- generate-design
PLACEHOLDER_RE = re.compile(r"\{\{(value|level|note|vl):([A-F]\.\d+\.\d+\.\d+)\}\}")


def cmd_generate_design(args) -> None:
    """テンプレートへDB値を機械代入して運用設計書Markdownを生成する。"""
    template_path = Path(args.template) if args.template else DEFAULT_TEMPLATE
    if not template_path.exists():
        die(f"テンプレートが見つかりません: {template_path}")
    conn = load_data(args.data)
    p = conn.execute("SELECT * FROM project").fetchone()
    items = {
        r["item_id"]: r
        for r in conn.execute(
            "SELECT i.item_id, i.duplicate_of, s.level, s.value, s.note"
            " FROM nfr_items i LEFT JOIN selections s ON s.item_id = i.item_id"
        )
    }

    text = template_path.read_text(encoding="utf-8")
    text = text.replace("{{system_name}}", p["system_name"])
    text = text.replace("{{model_system}}", str(p["model_system"]))
    text = text.replace("{{date}}", datetime.now(JST).strftime("%Y年%m月%d日"))

    unresolved = []

    def _resolve(m: re.Match) -> str:
        """プレースホルダをDB値に置換する。未登録は[要確認]を返す。"""
        kind, item_id = m.group(1), m.group(2)
        if item_id not in items:
            die(f"テンプレートの項目ID {item_id} がマスタに存在しません")
        row = items[item_id]
        value = (row["value"] or "").strip()
        level = (row["level"] or "").strip()
        note = (row["note"] or "").strip()
        if kind == "note":
            return note or "—"
        if kind == "level":
            return level or "—"
        if not value and not level:
            unresolved.append(item_id)
            return f"`[要確認: {item_id} 未登録]`"
        if kind == "vl":
            return f"{value}（{level}）" if value and level else (value or level)
        return value or level

    referenced = {m.group(2) for m in PLACEHOLDER_RE.finditer(text)}
    text = PLACEHOLDER_RE.sub(_resolve, text)

    # グレード外要求の対応表を機械生成
    fb = conn.execute(
        "SELECT * FROM feedback WHERE kind = 'out_of_grade' ORDER BY id"
    ).fetchall()
    if fb:
        lines = [
            "| # | 分類 | 要求内容 | 背景・理由 | 希望優先度 | 対応方針 | 状態 |",
            "|---|------|---------|-----------|-----------|---------|------|",
        ]
        for f in fb:
            lines.append(
                f"| {f['id']} | {f['classification'] or '—'} | {f['feedback']} |"
                f" {f['background'] or '—'} | {f['requested_priority'] or '—'} |"
                f" {f['response'] or '`[要確認: 対応方針未定]`'} |"
                f" {STATUS_LABELS.get(f['status'], f['status'])} |"
            )
        table = "\n".join(lines)
    else:
        table = "（グレード外要求はありません）"
    text = text.replace("{{out_of_grade_table}}", table)

    # カバレッジ検証: 全項目が本文（または重複代表項目）で参照されていること
    uncovered = [
        item_id for item_id, row in sorted(items.items())
        if item_id not in referenced and (row["duplicate_of"] or "") not in referenced
    ]

    Path(args.output).write_text(text, encoding="utf-8")
    registered = sum(1 for r in items.values() if (r["value"] or r["level"]))
    print(f"運用設計書を生成しました: {args.output}")
    print(f"  参照項目: {len(referenced)}項目 / 登録済み: {registered}/{len(items)}項目")
    print(f"  [要確認]（未登録）: {len(set(unresolved))}箇所")
    if uncovered:
        print(f"  警告: テンプレートで参照されていない項目: {', '.join(uncovered)}")
    else:
        print("  カバレッジ: 全項目が本文で参照されています（記載漏れなし）")
    conn.close()


# ---------------------------------------------------------------- export-word
INLINE_RE = re.compile(r"(\*\*.+?\*\*|`.+?`)")


def _add_runs(paragraph, text: str) -> None:
    """Markdownのインライン記法（太字・コード）をWordのrunに変換して追加する。"""
    for part in INLINE_RE.split(text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            paragraph.add_run(part[2:-2]).bold = True
        elif part.startswith("`") and part.endswith("`") and len(part) > 1:
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Consolas"
        else:
            paragraph.add_run(re.sub(r"\*([^*]+)\*", r"\1", part))


def _set_ja_font(doc, font_name: str = "游ゴシック") -> None:
    """ドキュメントの主要スタイルに日本語フォント（東アジア用含む）を設定する。"""
    from docx.oxml.ns import qn
    for style_name in (
        "Normal", "Title", "Heading 1", "Heading 2", "Heading 3", "Heading 4",
        "List Bullet", "List Number",
    ):
        try:
            style = doc.styles[style_name]
        except KeyError:
            continue
        style.font.name = font_name
        style.element.rPr.rFonts.set(qn("w:eastAsia"), font_name)


def _add_table(doc, rows: list) -> None:
    """Markdownの表行をWordの表（Table Grid・ヘッダー太字）へ変換する。"""
    cells = [
        [c.strip() for c in line.strip().strip("|").split("|")]
        for line in rows
        if not re.match(r"^\s*\|?[\s:|-]+\|?\s*$", line)
    ]
    if not cells:
        return
    ncols = max(len(r) for r in cells)
    table = doc.add_table(rows=len(cells), cols=ncols)
    table.style = "Table Grid"
    for ri, row in enumerate(cells):
        for ci in range(ncols):
            text = row[ci] if ci < len(row) else ""
            cell = table.cell(ri, ci)
            cell.text = ""
            para = cell.paragraphs[0]
            _add_runs(para, text)
            if ri == 0:
                for run in para.runs:
                    run.bold = True


def _markdown_to_docx(doc, md_text: str) -> None:
    """Markdownテキストを見出し・表・リスト・コードブロック対応でWordへ変換する。

    文書タイトル（先頭の唯一の ``#`` 見出し）は表紙と重複するためスキップし、
    以降の見出しを1段繰り上げる（``##`` の章 → Word Heading 1）。
    """
    lines = md_text.splitlines()
    heads = []
    for idx, raw in enumerate(lines):
        m = re.match(r"^(#{1,6})\s", raw.strip())
        if m:
            heads.append((idx, len(m.group(1))))
    skip_title_idx, level_offset = None, 0
    if heads and heads[0][1] == 1 and sum(1 for _, lv in heads if lv == 1) == 1:
        skip_title_idx = heads[0][0]
        level_offset = 1
    i, in_code = 0, False
    while i < len(lines):
        line = lines[i]
        if i == skip_title_idx:
            i += 1
            continue
        if line.strip().startswith("```"):
            in_code = not in_code
            i += 1
            continue
        if in_code:
            para = doc.add_paragraph()
            run = para.add_run(line)
            run.font.name = "Consolas"
            i += 1
            continue
        stripped = line.strip()
        if not stripped or re.fullmatch(r"-{3,}", stripped):
            i += 1
            continue
        m = re.match(r"^(#{1,6})\s+(.*)$", stripped)
        if m:
            level = min(max(len(m.group(1)) - level_offset, 1), 4)
            doc.add_heading(re.sub(r"[*`]", "", m.group(2)), level=level)
            i += 1
            continue
        if stripped.startswith("|"):
            block = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                block.append(lines[i])
                i += 1
            _add_table(doc, block)
            continue
        m = re.match(r"^\s*[-*]\s+(.*)$", line)
        if m:
            para = doc.add_paragraph(style="List Bullet")
            _add_runs(para, m.group(1))
            i += 1
            continue
        m = re.match(r"^\s*\d+\.\s+(.*)$", line)
        if m:
            para = doc.add_paragraph(style="List Number")
            _add_runs(para, m.group(1))
            i += 1
            continue
        if stripped.startswith(">"):
            para = doc.add_paragraph()
            run = para.add_run(stripped.lstrip("> "))
            run.italic = True
            i += 1
            continue
        para = doc.add_paragraph()
        _add_runs(para, stripped)
        i += 1


def cmd_export_word(args) -> None:
    """Markdown運用設計書をWordへ変換し、表紙と要件一覧・指摘対応表の付録を付与する。"""
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Pt
    except ImportError:
        die("python-docx がインストールされていません: pip install python-docx")
    if not Path(args.design).exists():
        die(f"運用設計書（Markdown）が見つかりません: {args.design}")

    conn = load_data(args.data)
    p = conn.execute("SELECT * FROM project").fetchone()
    items = _select_rows(conn)
    feedback = conn.execute(
        "SELECT f.*, i.item_name FROM feedback f"
        " LEFT JOIN nfr_items i ON i.item_id = f.item_id ORDER BY f.id"
    ).fetchall()

    doc = Document()
    _set_ja_font(doc)
    doc.styles["Normal"].font.size = Pt(10.5)

    # 表紙
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(f"\n\n\n{p['system_name']}\n運用設計書")
    run.font.size = Pt(28)
    run.bold = True
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.add_run(
        f"\nIPA非機能要求グレード2018準拠（モデルシステム{p['model_system']}）\n"
        f"{datetime.now(JST).strftime('%Y年%m月%d日')}"
    ).font.size = Pt(14)
    doc.add_page_break()

    # 本文（Markdown変換）
    _markdown_to_docx(doc, Path(args.design).read_text(encoding="utf-8"))

    # 付録: 非機能要件選択一覧（優先度別）
    doc.add_page_break()
    doc.add_heading("付録: 非機能要件選択一覧（優先度別）", level=1)
    doc.add_paragraph(
        "優先度は後続の設計・構築作業への影響度に基づく分類（高・中・低）です。"
    )
    for pr in ("高", "中", "低"):
        subset = [i for i in items if i["priority"] == pr]
        if not subset:
            continue
        doc.add_heading(f"優先度: {pr}（{len(subset)}項目）", level=2)
        table = doc.add_table(rows=1 + len(subset), cols=6)
        table.style = "Table Grid"
        for ci, h in enumerate(
            ["項目ID", "項目名", "選択レベル", "選択値・具体値", "備考", "顧客判定"]
        ):
            cell = table.cell(0, ci)
            cell.text = h
            cell.paragraphs[0].runs[0].bold = True
        for ri, it in enumerate(subset, start=1):
            vals = [
                it["item_id"], it["item_name"], it["level"] or "[要確認: 未登録]",
                it["value"] or "", it["note"] or "", it["customer_judgement"] or "未確認",
            ]
            for ci, v in enumerate(vals):
                table.cell(ri, ci).text = str(v)

    # 付録: 顧客指摘対応表
    if feedback:
        doc.add_heading("付録: 顧客指摘対応表", level=1)
        table = doc.add_table(rows=1 + len(feedback), cols=6)
        table.style = "Table Grid"
        for ci, h in enumerate(["#", "区分", "対象", "指摘・要求内容", "対応方針", "状態"]):
            cell = table.cell(0, ci)
            cell.text = h
            cell.paragraphs[0].runs[0].bold = True
        for ri, f in enumerate(feedback, start=1):
            target = (
                f"{f['item_id']} {f['item_name'] or ''}".strip()
                if f["kind"] == "item"
                else f"グレード外（{f['classification'] or '分類なし'}）"
            )
            vals = [
                str(f["id"]),
                "項目指摘" if f["kind"] == "item" else "グレード外要求",
                target, f["feedback"], f["response"] or "[要確認: 対応方針未定]",
                STATUS_LABELS.get(f["status"], f["status"]),
            ]
            for ci, v in enumerate(vals):
                table.cell(ri, ci).text = v

    doc.save(args.output)
    open_fb = sum(1 for f in feedback if f["status"] == "open")
    print(f"Word出力完了: {args.output}")
    print(f"  付録: 非機能要件{len(items)}項目 / 顧客指摘{len(feedback)}件（未対応 {open_fb}件）")
    conn.close()


# ---------------------------------------------------------------- main
def main() -> None:
    """サブコマンドを解析して各処理へディスパッチする。"""
    with contextlib.suppress(AttributeError, ValueError):
        signal.signal(signal.SIGPIPE, signal.SIG_DFL)
    parser = argparse.ArgumentParser(
        description="IPA非機能要求グレード ワークフロー管理CLI",
        formatter_class=argparse.RawDescriptionHelpFormatter,
    )
    sub = parser.add_subparsers(dest="command", required=True)

    p = sub.add_parser("init", help="データファイル（YAML）の初期化")
    p.add_argument("--data", required=True, help="データファイル（nfr.yaml）のパス")
    p.add_argument("--project", required=True, help="対象システム名")
    p.add_argument("--model-system", type=int, required=True, choices=[1, 2, 3])
    p.add_argument("--master", help="項目マスタCSV（省略時は同梱マスタ。指定時はデータファイルに記録される）")
    p.add_argument("--force", action="store_true", help="既存データファイルを上書き")
    p.set_defaults(func=cmd_init)

    p = sub.add_parser("validate", help="データファイルの全制約検証（手編集後の確認用）")
    p.add_argument("--data", required=True, help="データファイル（nfr.yaml）のパス")
    p.set_defaults(func=cmd_validate)

    p = sub.add_parser("hearing-sheet", help="ヒアリング質問一覧のMarkdown出力")
    p.add_argument("--data", required=True, help="データファイル（nfr.yaml）のパス")
    p.add_argument("--priority", choices=["高", "中", "低"], help="優先度で絞り込み")
    p.add_argument("--all", action="store_true", help="登録済み項目も含める（現在値つき）")
    p.add_argument("--output", help="出力先ファイル（省略時は標準出力）")
    p.set_defaults(func=cmd_hearing_sheet)

    p = sub.add_parser("register", help="ヒアリング結果のJSON登録")
    p.add_argument("--data", required=True, help="データファイル（nfr.yaml）のパス")
    p.add_argument("--input", required=True, help="JSONファイルパス（- で標準入力）")
    p.set_defaults(func=cmd_register)

    p = sub.add_parser("status", help="登録・指摘対応状況サマリー")
    p.add_argument("--data", required=True, help="データファイル（nfr.yaml）のパス")
    p.set_defaults(func=cmd_status)

    p = sub.add_parser("export-excel", help="顧客レビュー用Excel出力")
    p.add_argument("--data", required=True, help="データファイル（nfr.yaml）のパス")
    p.add_argument("--output", required=True, help="出力先 .xlsx パス")
    p.set_defaults(func=cmd_export_excel)

    p = sub.add_parser("import-feedback", help="顧客記入済みExcelの取り込み")
    p.add_argument("--data", required=True, help="データファイル（nfr.yaml）のパス")
    p.add_argument("--input", required=True, help="顧客記入済み .xlsx パス")
    p.set_defaults(func=cmd_import_feedback)

    p = sub.add_parser("list-feedback", help="指摘一覧（JSON）")
    p.add_argument("--data", required=True, help="データファイル（nfr.yaml）のパス")
    p.add_argument("--status", choices=["open", "accepted", "rejected", "reflected"])
    p.set_defaults(func=cmd_list_feedback)

    p = sub.add_parser("update-feedback", help="指摘の対応方針・状態更新")
    p.add_argument("--data", required=True, help="データファイル（nfr.yaml）のパス")
    p.add_argument("--id", type=int, required=True)
    p.add_argument("--response", help="対応方針")
    p.add_argument("--status", choices=["open", "accepted", "rejected", "reflected"])
    p.add_argument("--item-id", help="グレード外要求を項目に紐付ける場合の項目ID")
    p.set_defaults(func=cmd_update_feedback)

    p = sub.add_parser("check", help="要件値の整合性を機械検証（ERRORで終了コード1）")
    p.add_argument("--data", required=True, help="データファイル（nfr.yaml）のパス")
    p.set_defaults(func=cmd_check)

    p = sub.add_parser("generate-design", help="運用設計書Markdownの機械生成")
    p.add_argument("--data", required=True, help="データファイル（nfr.yaml）のパス")
    p.add_argument("--output", required=True, help="出力先 .md パス")
    p.add_argument("--template", help="テンプレート（省略時は同梱テンプレート）")
    p.set_defaults(func=cmd_generate_design)

    p = sub.add_parser("dump", help="DB内容の出力（設計書生成の入力）")
    p.add_argument("--data", required=True, help="データファイル（nfr.yaml）のパス")
    p.add_argument("--format", choices=["json", "markdown"], default="markdown")
    p.set_defaults(func=cmd_dump)

    p = sub.add_parser("export-word", help="Markdown運用設計書のWord変換＋付録追加")
    p.add_argument("--data", required=True, help="データファイル（nfr.yaml）のパス")
    p.add_argument("--design", required=True, help="運用設計書Markdownファイル")
    p.add_argument("--output", required=True, help="出力先 .docx パス")
    p.set_defaults(func=cmd_export_word)

    args = parser.parse_args()
    args.func(args)


if __name__ == "__main__":
    main()
